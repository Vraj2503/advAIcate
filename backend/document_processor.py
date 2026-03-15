"""
Document Processing Utilities
Handles file upload, text extraction, and chunking

SECURITY NOTES:
- Content-type verification is delegated to validators.py (single source of truth).
- Resource limits are enforced to prevent excessive page counts and memory exhaustion.
"""

import io
import logging
import PyPDF2
import docx
from typing import Dict, List, Optional
import tiktoken

from config import (
    MAX_PDF_PAGES,
    MAX_EXTRACTED_TEXT_LENGTH,
)
from validators import verify_content_matches_type, DocumentSecurityError

logger = logging.getLogger(__name__)


# ======================
# CUSTOM EXCEPTIONS
# ======================

class DocumentProcessingError(Exception):
    """Raised when document processing fails."""
    pass


class DocumentProcessor:
    """Process various document types and extract text"""

    def __init__(self):
        self.supported_types = {
            'application/pdf': self._extract_pdf_text,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': self._extract_docx_text,
            'text/plain': self._extract_text_file,
        }
        # Initialize tokenizer for chunking
        try:
            self.tokenizer = tiktoken.get_encoding("cl100k_base")
        except Exception as e:
            logger.warning("Failed to initialize tiktoken tokenizer: %s. Falling back to character-based chunking.", str(e))
            self.tokenizer = None

    # ==================== Text Extraction ====================

    def extract_text(self, file_content: bytes, file_type: str) -> str:
        """
        Extract text from a file based on its type.

        Args:
            file_content: Binary content of the file
            file_type: MIME type of the file

        Returns:
            Extracted text as string

        Raises:
            DocumentSecurityError: If content doesn't match claimed type
            DocumentProcessingError: If extraction fails
            ValueError: If file type is unsupported
        """
        extractor = self.supported_types.get(file_type)
        if not extractor:
            raise ValueError(f"Unsupported file type: {file_type}")

        # Single source of truth — validators.py handles all content verification
        verify_content_matches_type(file_content, file_type)

        logger.info("Extracting text from %s (%d bytes)", file_type, len(file_content))
        text = extractor(file_content)

        if len(text) > MAX_EXTRACTED_TEXT_LENGTH:
            logger.warning(
                "Extracted text truncated from %d to %d characters",
                len(text), MAX_EXTRACTED_TEXT_LENGTH
            )
            text = text[:MAX_EXTRACTED_TEXT_LENGTH]

        logger.info("Extracted %d characters of text", len(text))
        return text

    def _extract_pdf_text(self, file_content: bytes) -> str:
        """Extract text from PDF file with resource limits"""
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)

            # Enforce page limit
            page_count = len(pdf_reader.pages)
            if page_count > MAX_PDF_PAGES:
                logger.warning(
                    "PDF has %d pages, truncating to %d",
                    page_count, MAX_PDF_PAGES
                )
                page_count = MAX_PDF_PAGES

            text_parts = []
            for page_num in range(page_count):
                try:
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    if text and text.strip():
                        text_parts.append(f"[Page {page_num + 1}]\n{text}")
                except Exception as e:
                    logger.warning(
                        "Failed to extract text from PDF page %d: %s",
                        page_num + 1, str(e)
                    )
                    continue

            if not text_parts:
                raise DocumentProcessingError(
                    "No extractable text found in PDF. "
                    "The file may be scanned/image-based."
                )

            return "\n\n".join(text_parts)

        except DocumentProcessingError:
            raise
        except Exception as e:
            raise DocumentProcessingError(f"Error extracting PDF text: {str(e)}")

    def _extract_docx_text(self, file_content: bytes) -> str:
        """Extract text from DOCX file"""
        try:
            docx_file = io.BytesIO(file_content)
            doc = docx.Document(docx_file)

            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))

            if not text_parts:
                raise DocumentProcessingError(
                    "No extractable text found in DOCX."
                )

            return "\n\n".join(text_parts)

        except DocumentProcessingError:
            raise
        except Exception as e:
            raise DocumentProcessingError(f"Error extracting DOCX text: {str(e)}")

    def _extract_text_file(self, file_content: bytes) -> str:
        """Extract text from plain text file"""
        try:
            return file_content.decode('utf-8')
        except UnicodeDecodeError:
            try:
                return file_content.decode('latin-1')
            except UnicodeDecodeError:
                raise DocumentProcessingError(
                    "Failed to decode text file with UTF-8 or Latin-1 encoding"
                )

    # ==================== Text Chunking ====================

    def chunk_text(
        self,
        text: str,
        chunk_size: int = 500,
        overlap: int = 50
    ) -> List[Dict[str, any]]:
        """
        Split text into chunks with overlap for better context preservation.

        Args:
            text: Input text to chunk
            chunk_size: Maximum tokens per chunk
            overlap: Number of tokens to overlap between chunks

        Returns:
            List of dictionaries with chunk text and metadata
        """
        if not text or not text.strip():
            return []

        # Validate parameters
        chunk_size = max(50, min(chunk_size, 4000))
        overlap = max(0, min(overlap, chunk_size // 2))

        if not self.tokenizer:
            return self._chunk_by_characters(text, chunk_size * 4, overlap * 4)

        # Tokenize the text
        tokens = self.tokenizer.encode(text)
        chunks = []

        start = 0
        chunk_id = 0

        while start < len(tokens):
            end = start + chunk_size
            chunk_tokens = tokens[start:end]

            chunk_text = self.tokenizer.decode(chunk_tokens)

            chunks.append({
                'chunk_id': chunk_id,
                'text': chunk_text,
                'start_token': start,
                'end_token': min(end, len(tokens)),
                'token_count': len(chunk_tokens)
            })

            start += chunk_size - overlap
            chunk_id += 1

        logger.debug("Split %d tokens into %d chunks", len(tokens), len(chunks))
        return chunks

    def _chunk_by_characters(
        self,
        text: str,
        chunk_size: int,
        overlap: int
    ) -> List[Dict[str, any]]:
        """Fallback character-based chunking"""
        chunks = []
        start = 0
        chunk_id = 0

        while start < len(text):
            end = start + chunk_size
            chunk_text = text[start:end]

            chunks.append({
                'chunk_id': chunk_id,
                'text': chunk_text,
                'start_char': start,
                'end_char': min(end, len(text)),
                'char_count': len(chunk_text)
            })

            start += chunk_size - overlap
            chunk_id += 1

        return chunks

    # ==================== Utilities ====================

    def generate_document_summary(self, text: str, max_length: int = 500) -> str:
        """
        Generate a simple extractive summary of the document.

        Args:
            text: Full document text
            max_length: Maximum length of summary in characters

        Returns:
            Summary text
        """
        if not text:
            return ""

        if len(text) <= max_length:
            return text

        summary = text[:max_length]
        last_period = summary.rfind('.')
        last_newline = summary.rfind('\n')

        break_point = max(last_period, last_newline)
        if break_point > max_length * 0.7:
            summary = summary[:break_point + 1]

        return summary + "..."

    def is_supported(self, file_type: str) -> bool:
        """
        Check if file type is supported.

        NOTE: This only checks MIME type string. For full validation,
        use extract_text() which also verifies file content via magic bytes.
        """
        return file_type in self.supported_types


# ==================== Global Singleton ====================

_document_processor: Optional[DocumentProcessor] = None


def get_document_processor() -> DocumentProcessor:
    """Get or create the global document processor instance"""
    global _document_processor
    if _document_processor is None:
        _document_processor = DocumentProcessor()
    return _document_processor